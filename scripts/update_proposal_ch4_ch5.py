"""Update only findings-related text in Chapters Four and Five.

The script preserves all existing headings, figures, captions, and content
outside Chapters Four and Five.
"""
from __future__ import annotations

import shutil
from datetime import datetime
from pathlib import Path

from docx import Document


SOURCE = Path(r"C:\Users\USER\OneDrive\Desktop\Group 15 proposal.docx")
ALT_OUTPUT = Path(r"C:\Users\USER\OneDrive\Desktop\Group 15 proposal_UPDATED.docx")
REPO_OUTPUT = (
    Path(__file__).resolve().parents[1] / "Proposal" / "Group 15 proposal.docx"
)


REPLACEMENTS = {
    "The project successfully produced a predictive credit scoring prototype that uses supply chain transaction data to classify SME credit risk. The analytical pipeline transformed raw transaction records into structured variables suitable for machine learning, and the final model generated credit risk predictions that can support more objective lending decisions.": (
        "The project successfully produced and deployed a predictive credit scoring prototype that uses supply chain transaction data to classify SME credit risk. The analytical pipeline transforms uploaded or manually recorded transactions into structured behavioural variables. After an SME records at least five transactions, the system generates a credit score, risk band, creditworthy probability, explainable signals, and an indicative financing amount for SME and lender decision support."
    ),
    "The findings show that the model can be used as an alternative assessment tool for SMEs that lack sufficient collateral or formal financial records. This is important in the Tanzanian context, where traditional credit evaluation methods often exclude viable businesses.": (
        "The findings show that supply-chain transaction behaviour can support alternative assessment of SMEs that lack sufficient collateral or audited financial statements. This is important in the Tanzanian context, where viable informal and semi-formal businesses may be excluded by traditional evaluation. The deployed Ushirika portal operationalises this approach through separate bilingual workspaces for SMEs, lenders, administrators, and sub-administrators."
    ),
    "The preprocessing stage cleaned and prepared the transactional dataset by handling missing values, removing noise, detecting anomalies, and standardizing numerical features. Exploratory data analysis was used to examine patterns in payment behavior, transaction frequency, turnover, and default indicators, which improved the quality of the input data for modelling.": (
        "The preprocessing stage cleaned and prepared transaction data by validating required values, handling missing values, converting dates and amounts, normalising categorical values, and detecting unusual transactions. Numerical gaps were filled using median imputation, noisy continuous values were controlled using interquartile-range clipping, and the train/test process was stratified to preserve class balance. The platform accepts both English and Kiswahili CSV templates and manual transaction entry; counterparty TIN is optional because some informal buyers and suppliers do not possess one."
    ),
    "Feature engineering converted raw transaction activity into predictive variables that better reflect SME behavior. These included patterns related to payment consistency, transaction volume, account age, and compliance trends, making the dataset more informative for credit risk prediction.": (
        "Feature engineering converted raw transaction activity into predictive variables that reflect SME behaviour. These included payment reliability, average and maximum delay, transaction volume and frequency, account age, order completion, default and compliance rates, partner diversity, sales trend, on-time payment rate, intervals between transactions, buyer/supplier shares, and order-type mix. Rare high-value transactions were flagged so that a single unusual invoice would not inflate the indicative financing amount."
    ),
    "The first specific objective was to develop robust data preprocessing and feature engineering pipelines. This objective was achieved because the raw supply chain data was successfully transformed into a usable modelling dataset with cleaner, more meaningful variables.": (
        "The first specific objective was achieved. The implemented preprocessing and feature engineering pipeline transforms raw supply-chain transactions into seventeen behavioural predictors and outlier indicators. The same pipeline is used when SMEs upload a CSV file or record transactions manually, thereby connecting the analytical method directly to the working system."
    ),
    "The second specific objective was to compare machine learning ensemble techniques with classical regression models. This objective was addressed by implementing Logistic Regression as the baseline model and Random Forest as the ensemble model. The project framework indicates that ensemble methods were expected to capture complex non-linear relationships more effectively than classical regression.": (
        "The second specific objective was achieved by comparing Logistic Regression as the classical baseline with Random Forest as the ensemble model on the same stratified 80:20 train/test split. On the hold-out test data, Random Forest achieved 88.6% accuracy, 93.0% precision, 85.7% recall, 89.2% F1-score, and 95.8% ROC-AUC. Logistic Regression achieved 77.9% accuracy, 75.8% precision, 87.7% recall, 81.3% F1-score, and 87.5% ROC-AUC. Random Forest was therefore selected as the primary scorer because it provided the stronger overall discrimination and classification balance."
    ),
    "The third specific objective was to evaluate predictive reliability using standard metrics. This objective was achieved through the use of accuracy, precision, recall, F1-score, and ROC-AUC, which provided a balanced view of model performance and suitability for credit scoring decisions.": (
        "The third specific objective was achieved through hold-out accuracy, precision, recall, F1-score, ROC-AUC, and confusion-matrix evaluation. Five-fold GridSearchCV was applied on the training portion using ROC-AUC as the selection measure, while the final metrics were calculated on unseen test data. These controls, together with a minimum of five transactions before scoring and conservative score and financing rules, provide a transparent basis for interpreting the prototype's predictions."
    ),
    "Figure 4.5: Lender portfolio with SME risk classification, credit scores, and NIDA search.": (
        "Figure 4.5: Lender list of SMEs with risk classification, credit scores, and NIDA search."
    ),
    "The overall findings confirm that machine learning is suitable for alternative credit assessment in SME financing, especially where conventional financial data is limited. The preprocessing and feature engineering stage was crucial because the quality of the prediction depended heavily on the quality of the input data.": (
        "The overall findings support the suitability of machine learning for alternative SME credit assessment where conventional financial information is limited. The stronger Random Forest ROC-AUC of 95.8%, compared with 87.5% for Logistic Regression, indicates that non-linear relationships in payment behaviour, volume, delays, and value-chain roles were useful for ranking risk in the evaluated data. However, the result should be interpreted as prototype evidence rather than proof of performance across all Tanzanian SMEs."
    ),
    "The use of Random Forest is consistent with the literature cited in the proposal, which shows that ensemble models often outperform simpler regression approaches in credit scoring tasks. The project therefore contributes by showing how supply chain transaction data can be converted into a practical credit risk tool for Tanzanian SMEs.": (
        "The Random Forest result is consistent with Breiman (2001) and Lessmann et al. (2015), who demonstrate the value of ensemble methods for complex classification and credit-scoring tasks. The use of transaction behaviour as an alternative signal also agrees with Khandani et al. (2010), while the inclusion of buyer-supplier relationships reflects the value-chain financing perspective discussed by Klapper (2006). The project contributes by converting these ideas into a working Tanzanian SME credit-risk prototype rather than presenting only a theoretical model."
    ),
    "A notable limitation is that the project remains at prototype level rather than full operational deployment. In addition, the quality and completeness of transactional data may affect model reliability, especially where records are sparse, inconsistent, or incomplete.": (
        "Although the prototype is deployed online and its core workflows are operational, it is not yet a bank-production system. The training evidence still depends partly on synthetic bootstrap data and engineered labels because sufficiently large longitudinal repayment datasets from Tanzanian SMEs were not available. Transaction quality, sparse histories, free-tier cloud cold starts, and the absence of formal integration with national identity, banking, ERP, and credit-bureau systems may affect operational reliability and generalisability."
    ),
    "The study concludes that the proposed approach is useful for improving SME credit assessment because it shifts decision-making from collateral-heavy methods to data-driven analysis of supply chain behavior. The project also demonstrates that transaction data can provide valuable indicators of creditworthiness, especially for businesses operating in informal or semi-formal environments.": (
        "The study concludes that the proposed approach can improve SME credit assessment by supplementing collateral-heavy methods with data-driven analysis of supply-chain behaviour. The findings demonstrate that transaction data provides useful indicators of payment reliability, trading activity, delays, partner diversity, and order completion, especially for informal or semi-formal businesses with limited conventional records."
    ),
    "From the development process, one major lesson is that data preparation is as important as modelling. Another key conclusion is that ensemble learning methods are better suited for complex and non-linear transaction data than traditional regression models, making them more appropriate for this problem context.": (
        "The development process confirmed that data preparation is as important as modelling. Under the evaluated protocol, Random Forest performed better overall than Logistic Regression, particularly in accuracy, F1-score, and ROC-AUC. The conclusion is therefore limited to the evaluated dataset and protocol, but it supports the selection of Random Forest as the primary model for the current prototype."
    ),
    "The system is applicable as a prototype for financial institutions, fintech developers, and policymakers interested in promoting financial inclusion. However, its usefulness is still constrained by data availability, data quality, and the need for further testing in a real deployment environment.": (
        "The system is applicable as a deployed prototype for financial institutions, fintech developers, and policymakers interested in financial inclusion. Its bilingual SME, lender, administrator, and sub-administrator dashboards demonstrate the complete workflow from account management and transaction capture to scoring and lender review. Its usefulness nevertheless remains constrained by data availability, data quality, prototype hosting, and the need for supervised pilots using observed repayment outcomes."
    ),
    "Financial institutions should adopt data-driven credit assessment approaches that incorporate supply chain transaction behavior in addition to traditional financial records. This would help reduce exclusion of SMEs that are creditworthy but lack collateral or audited statements.": (
        "Financial institutions should pilot data-driven credit assessment that incorporates supply-chain transaction behaviour alongside, rather than immediately replacing, existing affordability, identity, and compliance checks. A phased pilot in one value-chain sector would allow lenders to compare Ushirika scores with observed repayment outcomes before wider adoption."
    ),
    "Authorities and lending organizations should invest in digital record-keeping systems and data-sharing infrastructure so that transaction data can be collected securely and used for fairer credit decisions. The model would also become more relevant if integrated with existing ERP systems, digital wallets, and merchant platforms.": (
        "Authorities and lending organisations should invest in secure digital record keeping, consent-based data sharing, and common transaction standards. Production implementation should integrate verified NIDA services, managed database infrastructure, ERP systems, digital wallets, merchant platforms, and auditable access controls while complying with applicable Tanzanian data-protection requirements."
    ),
    "For practical deployment, the system should be tested with larger and more diverse datasets from different SME sectors and regions in Tanzania. This would improve generalizability and reduce the risk of bias caused by limited or unbalanced data.": (
        "Before production use, the model should be validated on larger longitudinal datasets covering different SME sectors and regions of Tanzania. Evaluation should include fairness across groups, calibration, model drift, false-positive and false-negative costs, security testing, and independent human review of lending decisions."
    ),
    "Future research should focus on collecting more longitudinal transaction data and testing additional algorithms such as gradient boosting, XGBoost, or neural networks. Comparative studies could also be done to determine which model performs best under local Tanzanian SME conditions.": (
        "Future research should collect longitudinal transaction and observed repayment-outcome data, then compare Random Forest with gradient boosting, XGBoost, and other suitable algorithms under local Tanzanian SME conditions. Explainability methods such as SHAP and ongoing model-drift monitoring should also be evaluated before production adoption."
    ),
    "Another useful extension would be to develop a fully deployed web-based scoring platform with live dashboards, automated alerts, and model explainability features. This would make the system more transparent, user-friendly, and suitable for real financial decision-making.": (
        "Because a web-based scoring platform, role-based dashboards, bilingual forms, CSV import, and explainable score signals have already been implemented, future work should focus on production hardening rather than basic interface development. Priorities include managed PostgreSQL, verified identity services, multi-factor recovery, automated lender alerts, ERP and payment integrations, continuous model audit, and monitored pilots with partner financial institutions."
    ),
}


def update_document() -> Path:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup = SOURCE.with_name(f"Group 15 proposal_BACKUP_{timestamp}.docx")
    shutil.copy2(SOURCE, backup)

    doc = Document(str(SOURCE))
    original_shapes = len(doc.inline_shapes)
    replaced = set()

    in_scope = False
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text == "CHAPTER FOUR":
            in_scope = True
        elif text == "REFERENCES":
            in_scope = False
        if in_scope and text in REPLACEMENTS:
            paragraph.text = REPLACEMENTS[text]
            replaced.add(text)

    missing = set(REPLACEMENTS) - replaced
    if missing:
        raise RuntimeError(
            "Proposal text changed; replacements not applied:\n- "
            + "\n- ".join(sorted(missing))
        )

    target = SOURCE
    try:
        doc.save(str(target))
    except PermissionError:
        target = ALT_OUTPUT
        doc.save(str(target))

    check = Document(str(target))
    if len(check.inline_shapes) != original_shapes:
        raise RuntimeError("Figure count changed while updating proposal")

    REPO_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(target, REPO_OUTPUT)
    print(f"Updated {len(replaced)} Chapter 4-5 passages")
    print(f"Backup: {backup}")
    print(f"Saved: {target}")
    print(f"Repository copy: {REPO_OUTPUT}")
    return target


if __name__ == "__main__":
    update_document()
